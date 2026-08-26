#!/usr/bin/env python3
"""从需求文档中提取文字（及图片列表），生成分析报告，供设计测试点使用。

用法:
    python analyze_document.py <file> [-o OUTPUT_DIR] [-r REPORT.md]

支持:
    .txt/.md      标准库，直接读取
    .pdf          需要 pdfplumber (pip install pdfplumber)
    .docx         需要 python-docx (pip install python-docx)
    .png/.jpg     需要 pytesseract + Pillow 与本机 Tesseract OCR

缺少某格式的依赖时会给出安装提示，而不是崩溃。图片会被列在报告中，
请人工/多模态补充流程图、UI、状态图等关键图片的文字描述后再设计测试点。
"""
import argparse
import os
import sys


def read_text(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read(), []


def read_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        sys.exit("处理 PDF 需要 pdfplumber：pip install pdfplumber")
    text, images = [], []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text.append(page.extract_text() or "")
            for img in page.images:
                images.append(f"第 {i} 页图片 位置=({img.get('x0')},{img.get('top')})")
    return "\n".join(text), images


def read_docx(path):
    try:
        import docx
    except ImportError:
        sys.exit("处理 Word 需要 python-docx：pip install python-docx")
    d = docx.Document(path)
    text = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            text.append(" | ".join(c.text for c in row.cells))
    images = [r for r in d.part.rels.values() if "image" in r.reltype]
    return "\n".join(text), [f"内嵌图片: {r.target_ref}" for r in images]


def read_image(path):
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        sys.exit("图片 OCR 需要 pytesseract + Pillow，且本机安装 Tesseract OCR。")
    text = pytesseract.image_to_string(Image.open(path), lang="chi_sim+eng")
    return text, [f"输入图片: {os.path.basename(path)}"]


READERS = {
    ".txt": read_text, ".md": read_text,
    ".pdf": read_pdf, ".docx": read_docx,
    ".png": read_image, ".jpg": read_image, ".jpeg": read_image,
}


def main():
    ap = argparse.ArgumentParser(description="需求文档分析：提取文字与图片列表")
    ap.add_argument("file", help="需求文档路径")
    ap.add_argument("-o", "--output", default=".", help="输出目录")
    ap.add_argument("-r", "--report", help="分析报告路径 (.md)，默认 <输出目录>/analysis_report.md")
    args = ap.parse_args()

    if not os.path.isfile(args.file):
        sys.exit(f"文件不存在: {args.file}")
    ext = os.path.splitext(args.file)[1].lower()
    reader = READERS.get(ext)
    if reader is None:
        sys.exit(f"不支持的格式: {ext}（支持 {', '.join(sorted(READERS))}）")

    text, images = reader(args.file)
    os.makedirs(args.output, exist_ok=True)
    report = args.report or os.path.join(args.output, "analysis_report.md")

    lines = [
        f"# 需求文档分析报告\n",
        f"- 源文件: {args.file}",
        f"- 文字长度: {len(text)} 字符",
        f"- 图片数量: {len(images)}\n",
        "## 图片列表\n",
    ]
    lines += [f"- {d}" for d in images] or ["- （无）"]
    lines += ["\n## 文字内容\n", "```", text, "```"]
    with open(report, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"已生成分析报告: {report}（文字 {len(text)} 字符，图片 {len(images)} 个）")
    if images:
        print("提示: 请补充关键图片（流程图/UI/状态图）的文字描述后再设计测试点。")


if __name__ == "__main__":
    main()
